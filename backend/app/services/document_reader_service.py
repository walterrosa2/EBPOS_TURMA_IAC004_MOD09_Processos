import io
from docx import Document
from loguru import logger
from fastapi import HTTPException

class DocumentReadResult:
    def __init__(self, text: str, filename: str, size_bytes: int, num_paragraphs: int, num_characters: int, num_images: int):
        self.text = text
        self.filename = filename
        self.size_bytes = size_bytes
        self.num_paragraphs = num_paragraphs
        self.num_characters = num_characters
        self.num_images = num_images

    def to_dict(self) -> dict:
        return {
            "text": self.text,
            "metadata": {
                "filename": self.filename,
                "size_bytes": self.size_bytes,
                "num_paragraphs": self.num_paragraphs,
                "num_characters": self.num_characters,
                "num_images": self.num_images
            }
        }

class DocumentReaderService:
    @staticmethod
    def read_docx(file_bytes: bytes, filename: str) -> DocumentReadResult:
        """
        Lê um arquivo DOCX a partir de bytes, extrai parágrafos textuais e detecta imagens embutidas.
        """
        logger.info(f"Iniciando leitura do documento DOCX: {filename} ({len(file_bytes)} bytes)")
        
        try:
            # Carregar o DOCX a partir dos bytes na memória
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            
            paragraphs_text = []
            num_characters = 0
            
            # Extrair parágrafos operacionais e títulos
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    paragraphs_text.append(p_text)
                    num_characters += len(p_text)
            
            # Extrair tabelas e incluir no texto (proporciona maior riqueza analítica para a IA)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        # Unir as células da linha com separador
                        table_row_text = " | ".join(row_cells)
                        paragraphs_text.append(table_row_text)
                        num_characters += len(table_row_text)

            # Unir todo o texto extraído
            full_text = "\n".join(paragraphs_text)
            num_paragraphs = len(doc.paragraphs)
            
            # Detectar quantidade de imagens embutidas no documento
            num_images = 0
            try:
                for rel_id, rel in doc.part.related_parts.items():
                    if "image" in rel.content_type:
                        num_images += 1
            except Exception as e:
                logger.warning(f"Erro ao detectar imagens no DOCX (prosseguindo): {e}")

            logger.info(f"Leitura de {filename} finalizada: {num_paragraphs} parágrafos, {num_images} imagens, {num_characters} caracteres.")
            
            # Validar limite mínimo de texto extraído para garantir que o arquivo não está vazio ou puramente de imagens
            if num_characters < 50:
                logger.warning(f"Documento {filename} contém texto insuficiente ({num_characters} caracteres).")
                raise HTTPException(
                    status_code=400,
                    detail="O documento enviado possui conteúdo textual insuficiente para interpretação inteligente."
                )

            return DocumentReadResult(
                text=full_text,
                filename=filename,
                size_bytes=len(file_bytes),
                num_paragraphs=num_paragraphs,
                num_characters=num_characters,
                num_images=num_images
            )
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Erro ao processar arquivo DOCX {filename}: {str(e)}")
            raise HTTPException(
                status_code=400, 
                detail="Não foi possível ler o arquivo. Certifique-se de que é um documento .docx válido e não corrompido."
            )
