import os
import sys
import http.client

def smoke_test_import():
    # Caminhos absolutos
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    file_path = os.path.join(base_dir, "..", "prd_task_importacao_processo_ia", "sped fiscal exemplo.docx")
    file_path = os.path.abspath(file_path)
    
    temp_created = False
    if not os.path.exists(file_path):
        print("Aviso: Arquivo de amostra real não encontrado no repositório. Gerando um arquivo DOCX de teste operacional dinamicamente...")
        file_path = os.path.join(base_dir, "tests", "teste_importacao_temp.docx")
        file_path = os.path.abspath(file_path)
        
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Procedimento Operacional: Apuração Mensal de IPI', 0)
            doc.add_paragraph('Objetivo: Padronizar o processo de conferência e geração do imposto IPI.')
            doc.add_paragraph('Área: Fiscal')
            doc.add_paragraph('Etapa 1: Acessar o sistema Protheus e reprocessar saldos da apuração mensal.')
            doc.add_paragraph('Etapa 2: Validar o saldo devedor ou credor gerado com a contabilidade fiscal.')
            doc.add_paragraph('Etapa 3: Gerar a guia de recolhimento DARF via Sicalc Web caso haja imposto a pagar.')
            doc.add_paragraph('Regra: Se Cajamar possuir apuração zerada, então não gerar DARF.')
            doc.save(file_path)
            temp_created = True
            print(f"Arquivo temporário de teste gerado em: {file_path}")
        except Exception as e:
            print(f"Erro ao gerar DOCX de teste: {e}")
            sys.exit(1)
        
    print(f"Carregando arquivo para teste de fumaça: {file_path}")
    
    with open(file_path, "rb") as f:
        file_bytes = f.read()
        
    # Construindo o corpo multipart manualmente para usar biblioteca padrão sem dependência externa de requests
    boundary = "----WebKitFormBoundary7MA4YWxkTrZu0gW"
    content_type = f"multipart/form-data; boundary={boundary}"
    
    filename = os.path.basename(file_path)
    
    body = []
    body.append(f"--{boundary}".encode('utf-8'))
    body.append(f'Content-Disposition: form-data; name="file"; filename="{filename}"'.encode('utf-8'))
    body.append('Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document'.encode('utf-8'))
    body.append(b'')
    body.append(file_bytes)
    body.append(f"--{boundary}--".encode('utf-8'))
    body.append(b'')
    
    payload = b"\r\n".join(body)
    
    print("Enviando requisição POST para http://localhost:8000/api/processos/importar...")
    
    conn = http.client.HTTPConnection("localhost", 8000, timeout=120)  # timeout longo pois a IA processa o documento completo
    headers = {
        "Content-Type": content_type,
        "Content-Length": str(len(payload))
    }
    
    try:
        conn.request("POST", "/api/processos/importar", payload, headers)
        response = conn.getresponse()
        print(f"Status da resposta: {response.status} {response.reason}")
        
        response_data = response.read().decode('utf-8')
        print("Resposta do servidor:")
        print(response_data)
        
        if response.status == 201:
            print("🚀 SUCESSO! O serviço de importação inteligente respondeu perfeitamente.")
        else:
            print("❌ ERRO! O servidor respondeu com status de erro.")
            sys.exit(1)
            
    except Exception as e:
        print(f"Erro de conexão ou timeout: {e}")
        if temp_created and os.path.exists(file_path):
            os.remove(file_path)
        sys.exit(1)
    finally:
        conn.close()
        if temp_created and os.path.exists(file_path):
            try:
                os.remove(file_path)
                print("Arquivo temporário de teste removido da pasta de testes.")
            except Exception as e:
                print(f"Erro ao remover arquivo temporário: {e}")

if __name__ == "__main__":
    smoke_test_import()

